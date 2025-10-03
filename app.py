import streamlit as st
from docx import Document
import datetime
import os

# --- Replace placeholders robustly even in tables ---
def replace_placeholders(doc, replacements):
    import re
    pattern = re.compile(r"{{(.*?)}}")

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        for key, val in replacements.items():
            full_text = full_text.replace(f"{{{{{key}}}}}", str(val))
        paragraph.clear()
        paragraph.add_run(full_text)

    for p in doc.paragraphs:
        if pattern.search(p.text):
            replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if pattern.search(p.text):
                        replace_in_paragraph(p)

# --- Generate DOCX ---
def generate_docx(data, template_path="SALARY SLIP FORMAT.docx"):
    doc = Document(template_path)
    replace_placeholders(doc, data)
    file_name = f"salaryslip_{data['Name'].replace(' ', '_')}_{data['Month'].replace(' ', '_')}.docx"
    doc.save(file_name)
    return file_name

# --- Streamlit App ---
st.set_page_config(page_title="Single Salary Slip Generator", layout="centered")
st.title("📄 Individual Salary Slip Generator")

with st.form("salary_form"):
    st.subheader("Enter Employee Details")

    name = st.text_input("Name")
    designation = st.text_input("Designation")
    department = st.text_input("Department")
    total_days = st.number_input("Total Days", min_value=0, max_value=31)
    working_days = st.number_input("Working Days", min_value=0, max_value=31)
    weekly_off = st.number_input("Weekly Off", min_value=0, max_value=10)
    festival_off = st.number_input("Festival Off", min_value=0, max_value=10)
    paid_days = st.number_input("Paid Days", min_value=0, max_value=31)
    base = st.number_input("Base Salary", min_value=0.0)
    month = st.text_input("Month (e.g., July 2025)")

    salary = st.number_input("Salary", min_value=0.0)
    bonus = st.number_input("Bonus", min_value=0.0)
    other = st.number_input("Other", min_value=0.0)
    esi = st.number_input("ESI Deduction", min_value=0.0)
    advance_till = st.number_input("Advance Till Date", min_value=0.0)
    advance_deduct = st.number_input("Advance Deducted This Month", min_value=0.0)
    misc = st.number_input("MISC Deduction", min_value=0.0)

    submitted = st.form_submit_button("Generate Salary Slip")

if submitted:
    total = salary + bonus + other
    net_advance = advance_till - advance_deduct
    payable = total - (esi + advance_deduct + misc)
    payment_date = datetime.datetime.now().strftime("%d %B %Y")

    data = {
        "Name": name,
        "Designation": designation,
        "Department": department,
        "Total_Days": total_days,
        "Working_Days": working_days,
        "Weekly_Off": weekly_off,
        "Festival_Off": festival_off,
        "Paid_Days": paid_days,
        "Base": base,
        "Month": month,
        "Salary": salary,
        "Bonus": bonus,
        "Other": other,
        "Total": total,
        "ESI": esi,
        "Advance_Till_Date": advance_till,
        "Advance_Deduct": advance_deduct,
        "Net_Advance": net_advance,
        "MISC": misc,
        "Payable": payable,
        "Payment_Date": payment_date,
    }

    docx_path = generate_docx(data)

    with open(docx_path, "rb") as f:
        st.download_button(
            label="📄 Download Salary Slip",
            data=f,
            file_name=os.path.basename(docx_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
